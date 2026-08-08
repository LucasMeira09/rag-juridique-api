import requests
from bs4 import BeautifulSoup
from urllib.parse import urljoin
import os 
import re 
from io import BytesIO
from pypdf import PdfReader 
from docx import Document
from pathlib import Path

# Recupere les document des url passe et les transforme en texte netoyer et pret pour le pipeline

# Local config
_PROJECT_ROOT = Path(__file__).resolve().parent.parent.parent
DATA_DIR = os.getenv("DATA_DIR", str(_PROJECT_ROOT / "data"))

RAW_DIR = os.path.join(DATA_DIR, "raw_pdfs")
BEFORE_CLEAN_DIR = os.path.join(DATA_DIR, "before_clean_data")
CLEAN_DIR = os.path.join(DATA_DIR, "clean_data")


def list_files_local(directory_path):
    # Liste tous les fichiers dans un répertoire local
    try:
        if not os.path.isdir(directory_path):
            return []
        return [f for f in os.listdir(directory_path) if os.path.isfile(os.path.join(directory_path, f))]
    except Exception:
        return []

def read_text_local(file_path):
    # Lit un fichier texte local et retourne son contenu
    try:
        with open(file_path, "r", encoding="utf-8") as f:
            return f.read()
    except Exception as e:
        print(f"Erreur lors de la lecture: {e}")
        return None

def write_text_local(file_path, text):
    # Écrit un texte dans un fichier local
    try:
        os.makedirs(os.path.dirname(file_path), exist_ok=True)
        with open(file_path, "w", encoding="utf-8") as f:
            f.write(text)
        return True
    except Exception as e:
        print(f"Erreur lors de l'écriture: {e}")
        return False

def delete_file_local(file_path):
    # Supprime un fichier local
    try:
        os.remove(file_path)
        return True
    except Exception as e:
        print(f"Erreur lors de la suppression: {e}")
        return False

class TextScrapper():
    def __init__(self, url):
        
        try:
            self.headers = {
                # Utilisez un User-Agent commun pour Chrome, Firefox, ou autre
                'User-Agent': 'Mozilla/5.0 (Windows NT 10.0; Win64; x64) AppleWebKit/537.36 (KHTML, like Gecko) Chrome/100.0.4896.127 Safari/537.36'
            }
            self.url = url
            self.response = requests.get(self.url, self.headers)
            self.response.raise_for_status()

            self.html = self.response.text

            self.soup = BeautifulSoup(self.html, "html.parser")
        except:
            self.soup = ""

        self.pdf_urls = []    
        self.new_files_count = 0 # Compteur de nouveaux fichiers

        # Chemins locaux pour les différents dossiers
        self.raw_pdf = RAW_DIR
        self.output_folder = BEFORE_CLEAN_DIR
        self.final_folder = CLEAN_DIR
        
        # Créer les dossiers locaux si ils n'existent pas
        for directory in [self.raw_pdf, self.output_folder, self.final_folder]:
            os.makedirs(directory, exist_ok=True)

    def get_text(self):
        
        try:
            for a in self.soup.find_all("a", href=True, target="_blank"):
                href = a["href"]

                if href.lower() or "document" in str(href) or ".pdf" in str(href) or ".docx" in str(href) or "doc_num.php" in str(href):
                    pdf_url = urljoin(self.url, href)
                    self.pdf_urls.append(pdf_url)
        except Exception as e:
            print("Erreur de l'url: ", e)
        
    
    def download_text(self):
        self.get_text()

        file_list = list_files_local(self.raw_pdf)

        for pdf_url in self.pdf_urls:
            print("Has recovered :", pdf_url)
            try:
                request = requests.get(pdf_url, stream=True)
                request.raise_for_status()
            except Exception as e:
                print("Error web: ", e)
                continue

            # nom du fichier = dernière partie de l'URL
            if ".docx" in pdf_url or "doc_num.php" in pdf_url:

                # essayer de récupérer le vrai nom dans l'en-tête HTTP
                cd = request.headers.get("Content-Disposition")

                if cd:
                    # extraire le nom du fichier depuis l'en-tête
                    filename = re.findall('filename="?(.+)"?', cd)[0]
                else:
                    # fallback si pas d'en-tête → créer un nom propre
                    filename = pdf_url.split("/")[-1]
                    filename = filename.replace("?", "_").replace("=", "_")  # retirer caractères interdits
                    if not filename.lower().endswith(".docx"):
                        filename += ".docx"
            else:
                filename = pdf_url.split("/")[-1]
            
            filename = filename.strip().strip('"').strip("'")
            filename = re.sub(r'[<>:"/\\|?*]', '_', filename)
            filepath = os.path.join(self.raw_pdf, filename)

            try:
                if filename not in file_list:
                    # Télécharger le fichier en mémoire puis sauvegarder localement
                    file_data = b""
                    for chunk in request.iter_content(chunk_size=8192):
                        if chunk:
                            file_data += chunk
                    
                    # Sauvegarder sur disque local
                    with open(filepath, "wb") as f:
                        f.write(file_data)
                    print(f"→ Sauvegardé dans {filepath}")
                else:
                    print(f"Text: {filename}, already here.")
            except Exception as e:
                print("can't save:", e)

    def pdf_to_txt(self):
        
        # Convertit tous les PDFs d'un dossier en fichiers TXT localement
        
        # Récupérer tous les fichiers PDF depuis le dossier local
        pdf_files = list_files_local(self.raw_pdf)
        
        if not pdf_files:
            print(f"  Aucun fichier PDF trouvé dans le dossier '{self.raw_pdf}/'")
            print(f" Placez vos fichiers PDF dans le dossier '{self.raw_pdf}/' et relancez le script.")
            return
        
        print(f"\n{'='*80}")
        print(f"   CONVERSION PDF → TXT (Local) ".center(80))
        print(f"{'='*80}\n")
        print(f" Dossier source : {self.raw_pdf}/")
        print(f" Dossier destination : {self.output_folder}/")
        print(f" Nombre de PDFs trouvés : {len(pdf_files)}\n")
        print(f"{'='*80}\n")
        
        success_count = 0
        error_count = 0
        
        # liste avec tout les nom des fichier texte de before_clean_data
        text_list = list_files_local(self.output_folder)

        # Convertir chaque PDF
        for i, pdf_name in enumerate(pdf_files, 1):
            pdf_local_path = os.path.join(self.raw_pdf, pdf_name)
            root, extension = os.path.splitext(pdf_name)
            
            try:
                # Lire le fichier depuis le disque local
                with open(pdf_local_path, "rb") as f:
                    pdf_bytes = f.read()
                
                txt_name = pdf_name.replace(extension, '.txt')
                output_path = os.path.join(self.output_folder, txt_name)
                
                if txt_name not in text_list:
                    print(f"[{i}/{len(pdf_files)}]  Conversion de : {pdf_name}")
                    
                    text = ""
                    pages_count = 0
                    
                    # Essayer PDF d'abord
                    try:
                        reader = PdfReader(BytesIO(pdf_bytes))
                        pages_count = len(reader.pages)
                        # Extraire le texte de toutes les pages
                        for page_num, page in enumerate(reader.pages, 1):
                            page_text = page.extract_text() or ""
                            text += page_text
                        
                    except Exception:
                        # Si ce n'est pas un PDF, essayer DOCX
                        try:
                            doc = Document(BytesIO(pdf_bytes))
                            text = ""
                            for paragraph in doc.paragraphs:
                                text += paragraph.text + "\n"
                            pages_count = len(doc.paragraphs)  # Approximation
                        except Exception as e:
                            print(f"    Erreur: format non supporté ou fichier corrompu: {e}")
                            error_count += 1
                            continue
                    
                    # Écrire le texte localement
                    if write_text_local(output_path, text):
                        chars_count = len(text)
                        print(f"    Converti avec succès : {txt_name}")
                        print(f"    Pages : {pages_count} | Caractères : {chars_count:,}\n")
                        success_count += 1
                        self.new_files_count += 1 # Incrémenter le compteur global
                    else:
                        print(f"    Erreur lors de l'écriture du fichier\n")
                        error_count += 1
                else:
                    print(f"[{i}/{len(pdf_files)}]  {txt_name} existe déjà, ignoré.\n")
                    success_count += 1
                
            except Exception as e:
                print(f"  Erreur lors de la conversion : {str(e)}\n")
                error_count += 1
        
        # Résumé final
        print(f"{'='*80}")
        print(f"   RÉSUMÉ DE LA CONVERSION ".center(80))
        print(f"{'='*80}\n")
        print(f" Conversions réussies : {success_count}")
        print(f" Conversions échouées : {error_count}")
        print(f" Fichiers TXT disponibles dans : {self.output_folder}/\n")
        print(f"{'='*80}\n")


    def clean_text(self):
        text_path = list_files_local(self.output_folder)

        for text_name in text_path:
            text_directory = os.path.join(self.output_folder, text_name)

            # Lire le texte depuis le disque local
            text = read_text_local(text_directory)
            if text is None:
                print(f"Erreur lors de la lecture de {text_name}, ignoré.")
                continue

            # enleve les espace
            text_to_clean = re.sub(r'\s+', ' ', text)
            #enleve les espace devant la ponctuation double
            text_to_clean = re.sub(r'\s([:;?!])', r'\1', text_to_clean)
            # un espace apre ponctuation double
            text_to_clean = re.sub(r'([:;?!])([a-zA-Z0-9])', r'\1 \2', text_to_clean)
            # retire l espace devant la virgule et le point
            text_to_clean = re.sub(r'\s([,\.])', r'\1', text_to_clean)
            # reduit les espace autour des parenthese
            text_to_clean = re.sub(r'\s(\(|\))', r'\1', text_to_clean) # avent ( ou )
            text_to_clean = re.sub(r'(\()\s', r'\1', text_to_clean) # apre ( 
            text_to_clean = re.sub(r'\s(\))', r'\1', text_to_clean) # avant )
            # supprime les espace multiple 
            text_to_clean = re.sub(r'\s{2,}', ' ', text_to_clean)
            # suprime les /
            text_to_clean = text_to_clean.replace("/", "")
            # transmorme tout en minuscule
            text_to_clean = text_to_clean.lower()
            #
            text_to_clean = text_to_clean.strip()

            clean_text_directory = os.path.join(self.final_folder, text_name)

            # Écrire le texte nettoyé localement
            if not write_text_local(clean_text_directory, text_to_clean):
                print(f"Erreur lors de l'écriture de {text_name}")

    def clone_verifie(self):
        # verifie les si il y a des texte en double et les supprime

        text_list = list_files_local(self.final_folder)

        # recupere tout les texte
        for i, text_file in enumerate(text_list):
            try:
                text_directory = os.path.join(self.final_folder, text_file)
                text = read_text_local(text_directory)
                if text is None:
                    continue
            except:
                continue
            
            counter = 0
            for n in text_list:
                try:
                    text_directory_b = os.path.join(self.final_folder, n)
                    text_b = read_text_local(text_directory_b)
                    if text_b is None:
                        continue

                    if text == text_b and counter == 0:
                        counter += 1
                    elif text == text_b and counter > 0:
                        if delete_file_local(text_directory_b):
                            print(f"File {text_directory_b} removed")
                except:
                    pass

                
        
if __name__ == "__main__":
    # tout les site qu on veut scraper
    # peut avoir les url que vous vouler
    urls = [
        "https://environnement.brussels/pro/gestion-environnementale/gerer-les-dechets/parcours-dechets-professionnels-reduire-trier-et-gerer-vos-dechets-bruxelles"
    ]

    if len(urls) == 0:
        num = int(input("Combien de site a scraper?: "))
        
        for n in range(num):
            url = str(input("Votre url: "))
            urls.append(url)
    for u in urls:
        scrap = TextScrapper(u)
        # telecharge tout les texte
        scrap.download_text()
        scrap.pdf_to_txt()
        scrap.clean_text()

    scrap.clone_verifie()
